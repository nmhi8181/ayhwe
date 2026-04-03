from __future__ import annotations

import csv
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


BASE = Path(__file__).resolve().parent
OUT_DIR = BASE / "docx"
OUT_DIR.mkdir(exist_ok=True)
OUTPUT = OUT_DIR / "00_pakej_kajian_gabungan_formal_tesis.docx"


TITLE = (
    "HUBUNGAN ANTARA FAKTOR LINGUISTIK DAN KEBERKESANAN STRATEGI "
    "PENGAJARAN LITERASI BAHASA MELAYU MURID TAHAP 1"
)


APA_REFERENCES = [
    (
        "Abu Bakar, M. (2017). Kanak-kanak belajar membaca menggunakan buku cerita "
        "dwibahasa (Children learning to read with dual language storybooks). "
        "Jurnal Antarabangsa Alam dan Tamadun Melayu, 5(2). "
        "https://ejournal.ukm.my/jatma/article/view/73831"
    ),
    (
        "Chew, F. P. (2016). Masalah pembelajaran Bahasa Melayu dalam kalangan murid "
        "Cina sekolah rendah. Jurnal Pendidikan Bahasa Melayu, 6(2), 10-22. "
        "https://spaj.ukm.my/jpbm/index.php/jpbm/article/download/113/110"
    ),
    (
        "Chin, C. F., Chaiyanara, P. M., Bari, M. K., Li, L., & Zhang, D. (2012). "
        "Morphological awareness is important to word reading: Evidence from Malay "
        "literacy acquisition in Singapore primary school context. AARE Conference Papers. "
        "https://www.aare.net.au/publications/aare-conference-papers/show/6411/"
        "morphological-awareness-is-important-to-word-reading-evidence-from-malay-"
        "literacy-acquisition-in-singapore-primary-school-context"
    ),
    (
        "Hamzah, I. N. S., & Mahamod, Z. (2021). Strategi pengajaran dalam talian yang "
        "digunakan oleh guru Bahasa Melayu dalam meningkatkan kemahiran membaca murid "
        "sekolah rendah. Jurnal Pendidikan Bahasa Melayu, 11(2), 54-67. "
        "https://spaj.ukm.my/jpbm/index.php/jpbm/article/viewFile/267/202"
    ),
    (
        "Kementerian Pendidikan Malaysia. (2016). Dokumen Standard Kurikulum dan "
        "Pentaksiran Bahasa Melayu Tahun 1 Sekolah Kebangsaan. Bahagian Pembangunan "
        "Kurikulum. https://bukuteksdigital.my/wp-content/uploads/2020/06/"
        "DSKP-KSSR-Tahun-1-Bahasa-Melayu-SK.pdf"
    ),
    (
        "Kementerian Pendidikan Malaysia. (2022). Pelan Pembangunan Pendidikan Malaysia "
        "2013-2025: Laporan tahunan 2022. https://moe.gov.my/storage/files/shares/Dasar/"
        "PPPM/PPPM%20Laporan%20Tahunan%202022.pdf"
    ),
    (
        "Kementerian Pendidikan Malaysia. (2024). Manual pengoperasian pengesanan literasi "
        "dan numerasi serta program intervensi Tahun 1. "
        "https://cdn1.get-qr.com/files/66962f39df6a3bd4b1b27465/"
        "1724490801087-cc6eefa941d0f7f4f856665b9bdd528b.pdf"
    ),
    (
        "Low, J. Y., & Mahamod, Z. (2024). Strategi dan cabaran pengajaran kemahiran "
        "membaca yang digunakan oleh guru Bahasa Melayu dari Sekolah Jenis Kebangsaan "
        "Cina dalam meningkatkan kemahiran membaca murid Tahap 1. Jurnal Pendidikan "
        "Bahasa Melayu, 14(2). https://spaj.ukm.my/jpbm/index.php/jpbm/article/download/381/255"
    ),
    (
        "Low, J. Y., Wee, X. H., Khoo, Y. T., & Anuaruddin, N. F. F. (2024). Persepsi "
        "guru terhadap pengaruh bahasa ibunda semasa pembelajaran dan pengajaran Bahasa "
        "Melayu dalam kalangan murid di Sekolah Jenis Kebangsaan Cina. Jurnal Pendidikan "
        "Bahasa Melayu, 14(1). https://spaj.ukm.my/jpbm/index.php/jpbm/article/download/366/246"
    ),
    (
        "Nahar, N. (2020). Penguasaan kemahiran membaca dan menulis Bahasa Melayu dalam "
        "kalangan murid bukan penutur natif di Sekolah Jenis Kebangsaan. Issues in "
        "Language Studies, 9(1), 107-123. https://publisher.unimas.my/ojs/index.php/ILS/"
        "article/download/2223/987/6080"
    ),
    (
        "Sumardi, R. (2010). Phonological awareness and reading skills of Malay pre-school "
        "children aged 5 and 6 [Bachelor's thesis, Universiti Kebangsaan Malaysia]. "
        "https://mash.org.my/wp-content/uploads/Speech%202010%20Rozila%20Sumardi.pdf"
    ),
    (
        "U.S. Department of Education, Institute of Education Sciences, What Works "
        "Clearinghouse. (2017). Foundational skills to support reading for understanding "
        "in kindergarten through 3rd grade: Practice guide summary. "
        "https://ies.ed.gov/ncee/WWC/Docs/PracticeGuide/wwc_found_reading_summary_051517.pdf"
    ),
    (
        "Wan Ahmad, W. N. (2019). Amalan guru dalam melaksanakan kemahiran literasi "
        "Bahasa Melayu bagi Program LINUS di sekolah rendah. Jurnal Pendidikan Bahasa "
        "Melayu, 9(1), 1-11. https://spaj.ukm.my/jpbm/index.php/jpbm/article/download/187/154"
    ),
]


def set_a4_layout(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(4.0)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)


def configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)

    p_format = normal.paragraph_format
    p_format.line_spacing = 1.5
    p_format.first_line_indent = Cm(1.27)
    p_format.space_after = Pt(6)
    p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for style_name, size, bold in [
        ("Heading 1", 14, True),
        ("Heading 2", 12, True),
        ("Heading 3", 12, True),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = bold
        if style_name == "Heading 1":
            style.paragraph_format.space_before = Pt(12)
            style.paragraph_format.space_after = Pt(6)

    if "TOC Heading" not in doc.styles:
        toc_style = doc.styles.add_style("TOC Heading", WD_STYLE_TYPE.PARAGRAPH)
        toc_style.base_style = doc.styles["Heading 1"]


def add_update_fields_on_open(doc: Document):
    settings = doc.settings.element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        update = OxmlElement("w:updateFields")
        update.set(qn("w:val"), "true")
        settings.append(update)


def add_field(paragraph, instruction: str, placeholder: str = ""):
    run = paragraph.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)

    if placeholder:
        text = OxmlElement("w:t")
        text.text = placeholder
        run._r.append(text)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def set_page_numbering(section, start: int = 1, fmt: str = "decimal"):
    sect_pr = section._sectPr
    for child in list(sect_pr):
        if child.tag == qn("w:pgNumType"):
            sect_pr.remove(child)
    pg_num_type = OxmlElement("w:pgNumType")
    pg_num_type.set(qn("w:start"), str(start))
    pg_num_type.set(qn("w:fmt"), fmt)
    sect_pr.append(pg_num_type)


def add_page_number_to_footer(section):
    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(p, "PAGE")


def add_toc(doc: Document):
    p = doc.add_paragraph(style="Normal")
    add_field(
        p,
        r'TOC \o "1-3" \h \z \u',
        "Isi kandungan akan dikemas kini secara automatik apabila dokumen dibuka dalam Microsoft Word.",
    )


def clean_inline(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    text = text.replace("**", "")
    text = text.replace("*", "")
    text = text.replace("`", "")
    return text.strip()


def is_table_separator(line: str) -> bool:
    s = line.strip()
    if not (s.startswith("|") and s.endswith("|")):
        return False
    core = s.replace("|", "").replace("-", "").replace(":", "").replace(" ", "")
    return core == ""


def split_md_row(line: str):
    return [clean_inline(part.strip()) for part in line.strip().strip("|").split("|")]


def set_table_font(table, size=10):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(size)


def add_markdown(doc: Document, text: str, skip_first_heading: bool = False):
    lines = text.splitlines()
    i = 0
    first_heading_skipped = False
    in_code = False
    code_lines: list[str] = []

    def flush_code():
        nonlocal code_lines
        if code_lines:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.first_line_indent = Cm(0)
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(9)
            code_lines = []

    def is_plain_paragraph_start(value: str) -> bool:
        stripped = value.strip()
        if not stripped:
            return False
        if stripped.startswith("#"):
            return False
        if stripped.startswith("- "):
            return False
        if re.match(r"^\d+\.\s+", stripped):
            return False
        if stripped.startswith("```"):
            return False
        if stripped.startswith("|"):
            return False
        return True

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if in_code:
            if stripped.startswith("```"):
                flush_code()
                in_code = False
            else:
                code_lines.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("```"):
            in_code = True
            code_lines = []
            i += 1
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            if skip_first_heading and not first_heading_skipped:
                first_heading_skipped = True
                i += 1
                continue
            level = min(len(heading_match.group(1)), 3)
            doc.add_heading(clean_inline(heading_match.group(2)), level=level)
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            header = split_md_row(lines[i])
            rows = []
            j = i + 2
            while j < len(lines) and lines[j].strip().startswith("|"):
                rows.append(split_md_row(lines[j]))
                j += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for idx, cell in enumerate(header):
                table.rows[0].cells[idx].text = cell
            for row in rows:
                cells = table.add_row().cells
                for idx in range(len(header)):
                    cells[idx].text = row[idx] if idx < len(row) else ""
            set_table_font(table, size=10)
            i = j
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(clean_inline(stripped[2:]), style="List Bullet")
            p.paragraph_format.first_line_indent = Cm(0)
            i += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            item = re.sub(r"^\d+\.\s+", "", stripped)
            p = doc.add_paragraph(clean_inline(item), style="List Number")
            p.paragraph_format.first_line_indent = Cm(0)
            i += 1
            continue

        if is_plain_paragraph_start(line):
            parts = [stripped]
            j = i + 1
            while j < len(lines) and is_plain_paragraph_start(lines[j]):
                parts.append(lines[j].strip())
                j += 1
            doc.add_paragraph(clean_inline(" ".join(parts)))
            i = j
            continue

        i += 1

    if in_code:
        flush_code()


def add_csv_table(doc: Document, path: Path):
    with path.open("r", encoding="utf-8", newline="") as f:
        rows = list(csv.reader(f))
    header = rows[0]
    body = rows[1:]
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, cell in enumerate(header):
        table.rows[0].cells[idx].text = cell
    for row in body:
        cells = table.add_row().cells
        for idx in range(len(header)):
            cells[idx].text = row[idx] if idx < len(row) else ""
    set_table_font(table, size=9)


def add_cover_page(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    run = p.add_run("PAKEJ KAJIAN")
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run(TITLE)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(16)

    for text in [
        "Dokumen Sokongan Sorotan Literatur, Kerangka Konsep dan Matriks Sumber",
        "Disusun dalam format formal tesis untuk semakan dan suntingan lanjut",
        "",
        "Nama Penyelidik: ________________________________",
        "No. Matrik: ___________________________________",
        "Program: ______________________________________",
        "",
        "Tarikh: 3 April 2026",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if text:
            p.add_run(text)


def add_reference_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    p.add_run(text)


def build_document():
    doc = Document()
    doc.core_properties.title = TITLE
    doc.core_properties.subject = "Sorotan literatur dan kerangka konsep"
    doc.core_properties.language = "ms-MY"

    configure_styles(doc)
    add_update_fields_on_open(doc)

    cover_section = doc.sections[0]
    set_a4_layout(cover_section)
    add_cover_page(doc)

    toc_section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_a4_layout(toc_section)
    set_page_numbering(toc_section, start=1, fmt="lowerRoman")
    add_page_number_to_footer(toc_section)

    doc.add_heading("Isi Kandungan", level=1)
    add_toc(doc)

    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_a4_layout(body_section)
    set_page_numbering(body_section, start=1, fmt="decimal")
    add_page_number_to_footer(body_section)

    chapters = [
        ("Bab 1 Maklumat Dokumen", BASE / "README.md"),
        ("Bab 2 Sorotan Literatur", BASE / "01_sorotan_literatur.md"),
        ("Bab 3 Kerangka Konsep dan Cadangan Reka Bentuk Kajian", BASE / "02_kerangka_konsep_dan_reka_bentuk.md"),
    ]
    appendices = [
        ("Lampiran A Matriks Sumber", BASE / "03_matriks_sumber.csv"),
        ("Lampiran B Log Sumber dan Muat Turun", BASE / "04_log_sumber.md"),
    ]

    for idx, (chapter_title, path) in enumerate(chapters):
        if idx > 0:
            doc.add_page_break()
        doc.add_heading(chapter_title, level=1)
        add_markdown(doc, path.read_text(encoding="utf-8"), skip_first_heading=True)

    for appendix_title, path in appendices:
        doc.add_page_break()
        doc.add_heading(appendix_title, level=1)
        if path.suffix.lower() == ".csv":
            add_csv_table(doc, path)
        else:
            add_markdown(doc, path.read_text(encoding="utf-8"), skip_first_heading=True)

    doc.add_page_break()
    doc.add_heading("Rujukan", level=1)
    for reference in sorted(APA_REFERENCES):
        add_reference_paragraph(doc, reference)

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    output = build_document()
    print(output)
