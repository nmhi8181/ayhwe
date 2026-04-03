from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

from build_formal_docx import (
    add_markdown,
    add_page_number_to_footer,
    add_reference_paragraph,
    add_toc,
    add_update_fields_on_open,
    configure_styles,
    set_a4_layout,
    set_page_numbering,
)


BASE = Path(__file__).resolve().parent
SOURCE = BASE / "05_proposal_penuh_bab_1_3.md"
OUT_DIR = BASE / "docx"
OUT_DIR.mkdir(exist_ok=True)
OUTPUT = OUT_DIR / "05_proposal_penuh_bab_1_3_formal.docx"

TITLE = (
    "HUBUNGAN ANTARA FAKTOR LINGUISTIK DAN KEBERKESANAN STRATEGI "
    "PENGAJARAN LITERASI BAHASA MELAYU MURID TAHAP 1"
)


def add_cover_page(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(110)
    run = p.add_run("PROPOSAL KAJIAN")
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run(TITLE)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(15)

    for text in [
        "",
        "Disediakan dalam format formal untuk semakan penyelia dan pengembangan tesis",
        "",
        "Nama Penyelidik: ________________________________",
        "No. Matrik: ___________________________________",
        "Program: ______________________________________",
        "Fakulti / Universiti: ___________________________",
        "",
        "Tarikh: 3 April 2026",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if text:
            p.add_run(text)


def build_document():
    text = SOURCE.read_text(encoding="utf-8")
    body_text, references_text = text.split("\n## Rujukan\n", maxsplit=1)

    doc = Document()
    doc.core_properties.title = TITLE
    doc.core_properties.subject = "Proposal kajian penuh Bab 1 hingga Bab 3"
    doc.core_properties.language = "ms-MY"

    configure_styles(doc)
    add_update_fields_on_open(doc)

    cover_section = doc.sections[0]
    set_a4_layout(cover_section)
    add_cover_page(doc)

    toc_section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_a4_layout(toc_section)
    set_page_numbering(toc_section, start=1, fmt="lowerRoman")
    add_page_number_to_footer(toc_section)

    doc.add_heading("Isi Kandungan", level=1)
    add_toc(doc)

    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_a4_layout(body_section)
    set_page_numbering(body_section, start=1, fmt="decimal")
    add_page_number_to_footer(body_section)

    add_markdown(doc, body_text, skip_first_heading=True)

    doc.add_page_break()
    doc.add_heading("Rujukan", level=1)
    for line in references_text.splitlines():
        stripped = line.strip()
        if stripped:
            add_reference_paragraph(doc, stripped)

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    output = build_document()
    print(output)
